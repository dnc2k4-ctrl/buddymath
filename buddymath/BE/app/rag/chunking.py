"""
chunking.py – Document Chunker cho BuddyMath.
Tách tài liệu (PDF, DOCX, TXT, Markdown) thành chunks và đóng gói thành Document.
"""
from __future__ import annotations

import hashlib
import logging
import re
from pathlib import Path
from typing import Any

from app.rag.engine import Document

logger = logging.getLogger(__name__)

# ─── Defaults ────────────────────────────────────────────────────────────────
DEFAULT_CHUNK_SIZE    = 800
DEFAULT_CHUNK_OVERLAP = 150
MIN_CHUNK_SIZE        = 50


# ─── Regex patterns ──────────────────────────────────────────────────────────
_EXERCISE_START = re.compile(
    r"^(?:bài|câu|ví dụ|exercise|example|problem)\s*\d+",
    re.IGNORECASE | re.MULTILINE,
)
_SOLUTION_START = re.compile(
    r"^(?:lời giải|hướng dẫn|giải|solution|answer|hint)\s*[:.]?",
    re.IGNORECASE | re.MULTILINE,
)
_LATEX_BLOCK = re.compile(r"\$\$.*?\$\$|\\\[.*?\\\]", re.DOTALL)


# ─── Helpers ─────────────────────────────────────────────────────────────────
def _chunk_id(source_file: str, index: int, content: str) -> str:
    digest = hashlib.md5(content.encode()).hexdigest()[:8]
    stem   = Path(source_file).stem[:20] if source_file else "doc"
    return f"{stem}_{index:04d}_{digest}"


def _detect_doc_type(text: str) -> str:
    if _SOLUTION_START.search(text):
        return "solution"
    if _EXERCISE_START.search(text):
        return "exercise"
    return "theory"


def _protect_latex(text: str) -> tuple[str, dict[str, str]]:
    placeholders: dict[str, str] = {}
    def replacer(m):
        key = f"__LATEX_{len(placeholders):04d}__"
        placeholders[key] = m.group(0)
        return key
    return _LATEX_BLOCK.sub(replacer, text), placeholders


def _restore_latex(text: str, placeholders: dict[str, str]) -> str:
    for key, value in placeholders.items():
        text = text.replace(key, value)
    return text


# ─── Core Splitter ────────────────────────────────────────────────────────────
class RecursiveSplitter:
    SEPARATORS = ["\n\n", "\n", ". ", "! ", "? ", "; ", ", ", " ", ""]

    def __init__(self, chunk_size: int = DEFAULT_CHUNK_SIZE, overlap: int = DEFAULT_CHUNK_OVERLAP):
        self.chunk_size = chunk_size
        self.overlap    = overlap

    def split(self, text: str) -> list[str]:
        protected, placeholders = _protect_latex(text)
        raw_chunks = self._split_recursive(protected, self.SEPARATORS)
        return [
            _restore_latex(c, placeholders)
            for c in raw_chunks
            if len(c) >= MIN_CHUNK_SIZE
        ]

    def _split_recursive(self, text: str, separators: list[str]) -> list[str]:
        if len(text) <= self.chunk_size:
            return [text] if text.strip() else []

        sep       = separators[0] if separators else ""
        next_seps = separators[1:]

        if sep == "" or sep not in text:
            chunks, i = [], 0
            while i < len(text):
                end = min(i + self.chunk_size, len(text))
                chunks.append(text[i:end])
                i += self.chunk_size - self.overlap
            return chunks

        pieces  = text.split(sep)
        result: list[str] = []
        current = ""

        for piece in pieces:
            candidate = (current + sep + piece).lstrip(sep) if current else piece
            if len(candidate) <= self.chunk_size:
                current = candidate
            else:
                if current:
                    result.append(current)
                    overlap_text = current[-self.overlap:] if self.overlap else ""
                    current = (overlap_text + sep + piece).lstrip() if overlap_text else piece
                else:
                    result.extend(self._split_recursive(piece, next_seps))
                    current = ""

        if current:
            result.append(current)
        return result


# ─── Document Chunker ────────────────────────────────────────────────────────
class DocumentChunker:
    """Nhận file path, trả về list[Document]. Hỗ trợ .pdf .docx .txt .md .markdown .tex."""

    def __init__(self, chunk_size: int = DEFAULT_CHUNK_SIZE, overlap: int = DEFAULT_CHUNK_OVERLAP):
        self.splitter = RecursiveSplitter(chunk_size=chunk_size, overlap=overlap)

    # ── File readers ─────────────────────────────────────────────────────────
    def _read_pdf(self, path: Path) -> list[tuple[int, str]]:
        # Attempt 1: PyMuPDF (fitz)
        try:
            import fitz
            doc   = fitz.open(str(path))
            pages = [(i + 1, page.get_text("text")) for i, page in enumerate(doc)]
            doc.close()
            if sum(len(t) for _, t in pages) > 50:
                return pages
            logger.debug(f"PyMuPDF trích xuất ít text từ '{path.name}', thử pdfplumber…")
        except ImportError:
            pass

        # Attempt 2: pdfplumber
        try:
            import pdfplumber
            pages = []
            with pdfplumber.open(str(path)) as pdf:
                for i, page in enumerate(pdf.pages):
                    pages.append((i + 1, page.extract_text() or ""))
            if sum(len(t) for _, t in pages) > 50:
                return pages
        except ImportError:
            pass

        # Attempt 3: pdfminer
        try:
            from pdfminer.high_level import extract_pages
            from pdfminer.layout import LTTextContainer
            pages = []
            for i, layout in enumerate(extract_pages(str(path))):
                text = "\n".join(e.get_text() for e in layout if isinstance(e, LTTextContainer))
                pages.append((i + 1, text))
            if sum(len(t) for _, t in pages) > 50:
                return pages
        except ImportError:
            pass

        # Attempt 4: pymupdf4llm
        try:
            import pymupdf4llm
            md_text = pymupdf4llm.to_markdown(str(path))
            if md_text.strip():
                return [(1, md_text)]
        except ImportError:
            pass

        # Attempt 5: OCR (pytesseract + pdf2image)
        try:
            from pdf2image import convert_from_path
            import pytesseract
            logger.info(f"Thử OCR '{path.name}' bằng Tesseract…")
            pages_ocr: list[tuple[int, str]] = []
            for i, img in enumerate(convert_from_path(str(path), dpi=200)):
                pages_ocr.append((i + 1, pytesseract.image_to_string(img, lang="vie+eng")))
            if sum(len(t) for _, t in pages_ocr) > 50:
                logger.info(f"OCR thành công '{path.name}'")
                return pages_ocr
        except ImportError:
            logger.debug("pytesseract/pdf2image chưa cài – bỏ qua OCR.")
        except Exception as ocr_exc:
            logger.warning(f"OCR thất bại '{path.name}': {ocr_exc}")

        logger.warning(
            f"'{path.name}' là PDF dạng ảnh (scanned). Không trích xuất được text. "
            "Cài: pip install pdf2image pytesseract  và Tesseract OCR engine."
        )
        return [(1, "")]

    def _read_docx(self, path: Path) -> list[tuple[int, str]]:
        try:
            from docx import Document as DocxDocument
            doc        = DocxDocument(str(path))
            paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]

            pages: list[tuple[int, str]] = []
            current_page: list[str]      = []
            page_num                     = 1

            for para in doc.paragraphs:
                text = para.text.strip()
                if not text:
                    continue
                style = para.style.name.lower() if para.style else ""
                if style in ("heading 1", "heading 2") and current_page:
                    pages.append((page_num, "\n".join(current_page)))
                    page_num += 1
                    current_page = []
                current_page.append(text)

            if current_page:
                pages.append((page_num, "\n".join(current_page)))

            return pages if pages else [(1, "\n".join(paragraphs))]

        except ImportError:
            logger.warning("python-docx chưa cài. Đọc .docx dưới dạng raw text.")
            return [(1, path.read_text(encoding="utf-8", errors="replace"))]
        except Exception as exc:
            logger.error(f"Lỗi đọc docx '{path}': {exc}")
            return [(1, "")]

    def _read_text(self, path: Path) -> list[tuple[int, str]]:
        return [(1, path.read_text(encoding="utf-8", errors="replace"))]

    def _read_file(self, path: Path) -> list[tuple[int, str]]:
        suffix = path.suffix.lower()
        if suffix == ".pdf":
            return self._read_pdf(path)
        if suffix == ".docx":
            return self._read_docx(path)
        return self._read_text(path)

    # ── Public API ────────────────────────────────────────────────────────────
    def chunk_file(
        self,
        file_path: str,
        subject: str,
        topic: str,
        extra_metadata: dict[str, Any] | None = None,
    ) -> list[Document]:
        path = Path(file_path)
        if not path.exists():
            raise FileNotFoundError(f"Không tìm thấy file: {file_path}")

        pages          = self._read_file(path)
        extra_metadata = extra_metadata or {}
        documents: list[Document] = []
        chunk_index = 0

        for page_num, page_text in pages:
            if not page_text.strip():
                continue
            for raw in self.splitter.split(page_text):
                documents.append(
                    Document(
                        content=raw.strip(),
                        subject=subject,
                        topic=topic,
                        doc_type=_detect_doc_type(raw),
                        source_file=str(path),
                        page=page_num,
                        chunk_id=_chunk_id(str(path), chunk_index, raw),
                        extra_metadata=extra_metadata,
                    )
                )
                chunk_index += 1

        logger.info(
            f"Chunked '{path.name}': {len(pages)} pages → {len(documents)} chunks ({subject}/{topic})"
        )
        return documents

    def chunk_text(
        self,
        text: str,
        subject: str,
        topic: str,
        source_name: str                       = "inline",
        extra_metadata: dict[str, Any] | None = None,
    ) -> list[Document]:
        extra_metadata = extra_metadata or {}
        documents: list[Document] = []
        for i, raw in enumerate(self.splitter.split(text)):
            documents.append(
                Document(
                    content=raw.strip(),
                    subject=subject,
                    topic=topic,
                    doc_type=_detect_doc_type(raw),
                    source_file=source_name,
                    page=0,
                    chunk_id=_chunk_id(source_name, i, raw),
                    extra_metadata=extra_metadata,
                )
            )
        return documents
